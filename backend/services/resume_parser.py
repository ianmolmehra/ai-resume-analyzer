import io
import re
from pathlib import Path
from typing import Optional
from loguru import logger

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from schemas.resume import (
    ParsedResumeData, EducationItem, ExperienceItem, ProjectItem, CertificationItem
)
from utils.text_processor import (
    clean_text, extract_email, extract_phone, extract_linkedin,
    extract_github, extract_years_of_experience, calculate_date_range
)


EDUCATION_KEYWORDS = [
    "bachelor", "master", "phd", "doctorate", "b.sc", "m.sc", "b.tech", "m.tech",
    "b.e", "m.e", "mba", "b.a", "m.a", "associate", "diploma", "degree",
    "university", "college", "institute", "school", "graduated", "graduation"
]

EXPERIENCE_KEYWORDS = [
    "experience", "work history", "employment", "professional experience",
    "career", "positions held", "work experience"
]

PROJECT_KEYWORDS = ["projects", "personal projects", "academic projects", "portfolio"]

CERTIFICATION_KEYWORDS = ["certifications", "certificates", "credentials", "licenses", "courses"]

SKILL_KEYWORDS = ["skills", "technical skills", "core competencies", "expertise", "technologies"]

EDUCATION_DEGREES = {
    "phd": "PhD", "doctorate": "PhD", "master": "Master's",
    "msc": "Master's", "m.sc": "Master's", "mba": "MBA",
    "m.tech": "M.Tech", "m.e": "M.E",
    "bachelor": "Bachelor's", "bsc": "Bachelor's", "b.sc": "Bachelor's",
    "b.tech": "B.Tech", "b.e": "B.E", "b.a": "BA",
    "associate": "Associate", "diploma": "Diploma"
}


def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return clean_text(text)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, falling back to PyPDF2")

    if HAS_PYPDF2:
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"PyPDF2 failed: {e}")

    return clean_text(text)


def extract_text_from_docx(file_path: str) -> str:
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed")
    doc = DocxDocument(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return clean_text("\n".join(paragraphs))


def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported file type: {ext}")


def split_into_sections(text: str) -> dict:
    """Split resume text into logical sections."""
    lines = text.split('\n')
    sections = {}
    current_section = "header"
    current_content = []

    section_map = {
        "education": EDUCATION_KEYWORDS[:6],
        "experience": EXPERIENCE_KEYWORDS,
        "projects": PROJECT_KEYWORDS,
        "certifications": CERTIFICATION_KEYWORDS,
        "skills": SKILL_KEYWORDS,
        "summary": ["summary", "objective", "profile", "about me"],
    }

    for line in lines:
        line_lower = line.strip().lower()
        matched_section = None
        for section, keywords in section_map.items():
            if any(kw in line_lower for kw in keywords) and len(line.strip()) < 60:
                matched_section = section
                break

        if matched_section:
            sections[current_section] = "\n".join(current_content)
            current_section = matched_section
            current_content = []
        else:
            current_content.append(line)

    sections[current_section] = "\n".join(current_content)
    return sections


def extract_name(text: str) -> Optional[str]:
    """Try to extract name from the first few lines."""
    lines = [l.strip() for l in text.split('\n') if l.strip()][:8]
    for line in lines:
        # Skip lines that look like contact info
        if any(c in line for c in ['@', 'http', '+', '|', '/', '\\']) and len(line) > 40:
            continue
        if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}$', line):
            return line
        if re.match(r'^[A-Z]{2,}(?:\s+[A-Z]{2,})+$', line):
            return line.title()
    return None


def parse_education(text: str) -> list:
    education = []
    # Split into chunks by double newlines or common separators
    chunks = re.split(r'\n{2,}|(?=\b(?:Bachelor|Master|PhD|B\.Tech|M\.Tech|B\.E|M\.E|MBA|B\.Sc|M\.Sc)\b)', text)

    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk or len(chunk) < 10:
            continue

        degree = None
        for key, val in EDUCATION_DEGREES.items():
            if re.search(rf'\b{re.escape(key)}\b', chunk, re.IGNORECASE):
                degree = val
                break

        if not degree and not any(kw in chunk.lower() for kw in ["university", "college", "institute", "school"]):
            continue

        year_match = re.search(r'\b(19|20)\d{2}\b', chunk)
        year = year_match.group() if year_match else None

        uni_match = re.search(
            r'(?:University|College|Institute|School|Academy)\s+(?:of\s+)?[A-Z][^\n,]*',
            chunk, re.IGNORECASE
        )
        university = uni_match.group().strip() if uni_match else None

        if not university:
            lines = chunk.split('\n')
            for line in lines:
                if any(kw in line.lower() for kw in ["university", "college", "institute"]):
                    university = line.strip()
                    break

        education.append(EducationItem(
            degree=degree,
            university=university,
            graduation_year=year,
        ))

    return education[:4]  # cap at 4


def parse_experience(text: str) -> list:
    experience = []
    # Split by company/role patterns
    chunks = re.split(r'\n(?=[A-Z][^\n]{5,60}\n)', text)

    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk or len(chunk) < 20:
            continue

        lines = chunk.split('\n')
        role = None
        company = None
        duration = None

        # First line often is role or company
        if lines:
            first = lines[0].strip()
            if len(first) < 80:
                role = first

        # Look for duration patterns
        duration_match = re.search(
            r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|June|July|August|September|October|November|December)\.?\s+\d{4})\s*[-–—to]+\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|June|July|August|September|October|November|December)\.?\s+\d{4}|[Pp]resent|[Cc]urrent)',
            chunk, re.IGNORECASE
        )
        if duration_match:
            start = duration_match.group(1)
            end = duration_match.group(2)
            duration = f"{start} - {end}"
            years = calculate_date_range(start, end)
        else:
            years = 0.0

        if role:
            experience.append(ExperienceItem(
                role=role,
                company=company,
                duration=duration,
                description="\n".join(lines[1:])[:500] if len(lines) > 1 else None,
            ))

    return experience[:8]


def parse_projects(text: str) -> list:
    projects = []
    chunks = re.split(r'\n(?=[A-Z•\-\*])', text)

    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk or len(chunk) < 15:
            continue

        lines = chunk.split('\n')
        name = lines[0].strip().lstrip('•-* ').strip()
        desc = " ".join(lines[1:])[:300] if len(lines) > 1 else ""

        # Extract tech stack from parentheses or after colon
        tech_match = re.findall(r'\b(?:Python|JavaScript|React|Node|Django|Flask|SQL|MongoDB|AWS|Docker|Git|Java|C\+\+|TypeScript|Vue|Angular|PostgreSQL|MySQL|Redis|FastAPI|TensorFlow|PyTorch|Scikit-learn|Pandas|NumPy)\b', chunk, re.IGNORECASE)
        tech_stack = list(set(tech_match))

        projects.append(ProjectItem(
            name=name[:100],
            description=desc,
            tech_stack=tech_stack,
        ))

    return projects[:8]


def parse_certifications(text: str) -> list:
    certifications = []
    lines = [l.strip() for l in text.split('\n') if l.strip() and len(l.strip()) > 5]

    for line in lines:
        year_match = re.search(r'\b(20\d{2}|19\d{2})\b', line)
        year = year_match.group() if year_match else None
        name = re.sub(r'\b(20\d{2}|19\d{2})\b', '', line).strip().lstrip('•-* ')
        if name and len(name) > 3:
            certifications.append(CertificationItem(
                name=name[:200],
                year=year,
            ))

    return certifications[:10]


def detect_highest_education(education_list: list) -> Optional[str]:
    order = ["PhD", "Master's", "MBA", "M.Tech", "M.E", "Bachelor's", "B.Tech", "B.E", "BA", "Associate", "Diploma"]
    if not education_list:
        return None
    for degree in order:
        for edu in education_list:
            if edu.degree and degree.lower() in edu.degree.lower():
                return edu.degree
    return education_list[0].degree if education_list else None


def parse_resume(file_path: str) -> ParsedResumeData:
    logger.info(f"Parsing resume: {file_path}")
    raw_text = extract_text(file_path)

    if not raw_text or len(raw_text) < 50:
        raise ValueError("Could not extract text from resume. File may be empty or image-based.")

    sections = split_into_sections(raw_text)
    header_text = sections.get("header", raw_text[:2000])

    full_name = extract_name(header_text) or extract_name(raw_text[:500])
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    linkedin = extract_linkedin(raw_text)
    github = extract_github(raw_text)

    education = parse_education(sections.get("education", ""))
    experience = parse_experience(sections.get("experience", ""))
    projects = parse_projects(sections.get("projects", ""))
    certifications = parse_certifications(sections.get("certifications", ""))

    total_exp = extract_years_of_experience(raw_text)

    return ParsedResumeData(
        full_name=full_name,
        email=email,
        phone=phone,
        linkedin_url=linkedin,
        github_url=github,
        education=education,
        experience=experience,
        projects=projects,
        certifications=certifications,
        total_experience_years=total_exp,
        highest_education=detect_highest_education(education),
        skills=[],  # will be populated by skill extractor
    )
